# -*- coding: utf-8 -*-

from collections import defaultdict
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import lxml
import os
import requests
import json
import qrcode
import hashlib
import re
import datetime
import logging
import zipfile
from threading import Thread
from . import config


config = config.config()
treport = config['treport']


def tree():
    return defaultdict(tree)


def add_run_self(para, text, font_name=None, font_size=9, font_color=(89, 89, 89), font_italic=False, alignment='center'):
    """段落添加文字，可修改字体，大小，颜色，斜体, 对齐方式"""
    run = para.add_run(str(text))
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    a, b, c = font_color
    font_color = RGBColor(a, b, c)
    run.font.color.rgb = font_color
    run.font.italic = font_italic
    if alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if alignment == 'left':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if alignment == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if alignment == 'other':
        pass


def add_tables_self(table, row_num, merge_col_list=None):
    """在表格后追加行，合并所需列"""
    if merge_col_list is None:
        merge_col_list = []
    table_row_num = len(table.rows)
    for i in range(row_num):
        table.add_row()
    if not merge_col_list:
        return
    for col in merge_col_list:
        row = table_row_num + row_num - 1
        table.cell(table_row_num, col).merge(table.cell(row, col))


def adjust_table_width(table, widthlist):
    """调整表格宽度"""
    colnum = len(widthlist)
    for i in range(colnum):
        for cell in table.columns[i].cells:
            cell.width = Cm(widthlist[i])


def find_para(tpl, expect):
    """
    根据内容找位置
    :param tpl:
    :param expect:
    :return:
    """
    for index, paragraph in enumerate(tpl.paragraphs):
        if paragraph.text == expect:
            return index


def set_updatefields_true(doc):
    """更新word目录"""
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    # add child to doc.settings element
    element_updatefields = lxml.etree.SubElement(doc.settings.element, namespace + "updateFields")
    element_updatefields.set(namespace + "val", "true")


def cms(sample):
    """依据样本编号获取患者信息"""
    useragent = 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_6)\
     AppleWebKit/537.36 (KHTML, like Gecko)\
      Chrome/74.0.3729.169 Safari/537.36'
    header = {'User-Agent': useragent}
    posturl = 'http://cms.topgen.com.cn/user/login/auth'
    # postData = {"userName": "xiecaixia", "password": "xcx50800383"}
    postdata = {"userName": "kaitian", "password": "top123"}
    response = requests.post(posturl, params=postdata, headers=header)
    text = response.text
    if response.status_code == 200:
        token = json.loads(text)["data"]["accessToken"]
        # print(token)
        payload = {"search[sampleSn]": sample, "accessToken": token}
        r = requests.get("http://cms.topgen.com.cn/sample/sample/search?report=true", params=payload)
        # print(r)
    else:
        raise Exception("网站不能在此账号下访问!", response.status_code)
    info_list = json.loads(r.text)['data'][0]
    if info_list['gender'] == '男':
        info_list['gender1'] = info_list['name'] + '先生'
    else:
        info_list['gender1'] = info_list['name'] + '女士'
    if info_list['agentName'].find('不体现') > -1:
        info_list['agentName'] = 'NA'
    info_list['makeTime1'] = datatran(info_list['makeTime'])
    info_list['receiveTime1'] = datatran(info_list['receiveTime'])
    info_list['decTime'] = datatran(info_list['receiveTime'], flag=1, add=1)
    for tm in info_list:
        if tm == 'decTime':
            continue
        if info_list[tm] == '':
            info_list[tm] = 'NA'
        info_list[tm] = re.split(r'[\s(]', str(info_list[tm]))[0]
    return info_list


def qrcodeself(path, sample):
    """依据样品编号生成二维码，指定路径"""
    hashself = hashlib.md5(sample.encode()).hexdigest()
    if treport:
        url = 'http://treport.topgen.com.cn/' + hashself + '.html'
    else:
        url = 'http://report.topgen.com.cn/' + hashself + '.html'
    img = qrcode.make(url)
    img.save(os.path.join(path, sample + ".png"))


def datatran(datainfo=None, flag=0, add=0):
    """
    时间格式转换
    :param datainfo: 时间可以是脚本运行时间，或者提供时间格式化
    :param flag: 是否带时和秒
    :param add: 增加一小时
    :return:
    """
    if not datainfo:
        datainfo = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    datainfo = str(datainfo)
    year, month, day, hour, minute, sec = re.split(r'[-\s+:]', datainfo)
    if flag:
        if add:
            hour = int(hour) + 1
        datanew = '{}年{}月{}日 {}:{}'.format(year, month.lstrip('0'), day.lstrip('0'), hour, minute)
    else:
        datanew = '{}年{}月{}日'.format(year, month.lstrip('0'), day.lstrip('0'))
    return datanew


def reportlog(name):
    """
    日志
    :param name:
    :return: 返回logger对象
    """
    basedir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'result', 'log')
    logger = logging.getLogger('main')
    logpath = os.path.join(basedir, name+'.log.txt')
    logfile = logging.FileHandler(logpath, mode='w')
    logfomat = logging.Formatter('%(message)s\t%(asctime)s')
    logfile.setFormatter(logfomat)
    logger.addHandler(logfile)
    return logpath


def filetozip(filelist, xlsxname):
    """
    压缩文件
    :param xlsxname:
    :param filelist:
    :return:
    """
    if len(filelist) == 0:
        return None
    path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'result', 'zip', xlsxname+'.zip')
    myzip = zipfile.ZipFile(path, "w")
    for fp in filelist:
        fn = os.path.basename(fp)
        myzip.write(fp, fn)  # 第一个参数为该文件的全路径；第二个参数为文件在zip中的相对路径
    myzip.close()
    return path


def drugcmp(drug1, drug2):
    """
    特定药物排序
    :param drug1:
    :param drug2:
    :return:
    """
    durglevel = {
        '奥希替尼': 1,
        '厄洛替尼': 2,
        '吉非替尼': 3,
        '埃克替尼': 4,
        '阿法替尼': 5,
        '达可替尼': 6,
        '厄洛替尼+Ramucirumab[雷莫芦单抗]': 7,
        '厄洛替尼+贝伐珠单抗': 8,
    }
    if drug1 in durglevel and drug2 in durglevel:
        drug1level = durglevel[drug1]
        drug2level = durglevel[drug2]
        if drug1level < drug2level:
            return -1
        else:
            return 1
    elif drug1 in durglevel and drug2 not in durglevel:
        return -1
    elif drug1 not in durglevel and drug2 in durglevel:
        return 1
    else:
        if drug1 > drug2:
            return -1
        else:
            return 1


class MyThread(Thread):
    def __init__(self, func, args):
        super(MyThread, self).__init__()
        self.func = func
        self.args = args

    def run(self):
        self.result = self.func(*self.args)

    def get_result(self):
        try:
            return self.result
        except Exception:
            return None


rootpath = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))
newline = '</w:t><w:br/><w:t xml:space="preserve">'
newpara = '</w:t></w:r></w:p><w:p><w:r><w:t xml:space="preserve">'